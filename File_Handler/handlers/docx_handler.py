# handlers/docx_handler.py
#
# This module handles the processing of DOCX files uploaded to a Discord channel.
# It provides options via buttons to either create a thread with DOCX pages as images
# or post the DOCX content in the current channel.
#
# The module uses python-docx for parsing DOCX files and Pillow (PIL) for creating images from text.

# Import necessary libraries
import discord  # Main Discord API library
from discord.ui import Button, View  # UI components for interactive buttons
import docx  # Python-docx for reading DOCX files
from io import BytesIO  # For handling binary data in memory
import asyncio  # For asynchronous operations
from discord.errors import Forbidden  # For handling permission errors
from PIL import Image, ImageDraw, ImageFont  # Pillow for image creation
import textwrap  # For wrapping text in images
import os  # For file path operations

async def convert_docx_to_images(docx_bytes, max_width=1024, font_size=16, margin=20):
    """Convert DOCX bytes to a list of image streams.
    
    This function takes the binary content of a DOCX file and converts it to
    a series of images, with each image representing a "page" of content.
    
    Args:
        docx_bytes: The binary content of the DOCX file
        max_width: Maximum width of the generated images in pixels (default: 1024)
        font_size: Font size to use for text (default: 16)
        margin: Margin around text in pixels (default: 20)
        
    Returns:
        A list of BytesIO objects containing PNG images of the document content
    """
    try:
        # Save the docx bytes to a temporary file since python-docx can't read from BytesIO directly
        temp_file = BytesIO(docx_bytes)
        
        # Open the document using python-docx
        doc = docx.Document(temp_file)
        
        # Get all paragraphs from the document, filtering out empty ones
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Get all tables from the document and convert to text
        table_texts = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                # Join cell text with pipe separators for visual clarity
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():  # Skip empty rows
                    table_text.append(row_text)
            if table_text:  # Only add non-empty tables
                table_texts.append("\n".join(table_text))
        
        # Combine paragraphs and table texts in document order
        all_text = []
        para_index = 0
        table_index = 0
        
        # A simplistic approach to maintain document order
        # In a real-world scenario, you would need more sophisticated logic
        for item in doc.element.body:
            if item.tag.endswith('p') and para_index < len(paragraphs):
                # This is a paragraph
                all_text.append(paragraphs[para_index])
                para_index += 1
            elif item.tag.endswith('tbl') and table_index < len(table_texts):
                # This is a table
                all_text.append("TABLE START")
                all_text.append(table_texts[table_index])
                all_text.append("TABLE END")
                table_index += 1
        
        # If there's nothing in the document, return empty list
        if not all_text:
            return []
        
        # Split the content into pages (approximately 3000 characters per page)
        # This helps ensure each image isn't too large
        pages = []
        current_page = []
        current_length = 0
        
        for item in all_text:
            # If adding this item would exceed our target page size, start a new page
            if current_length + len(item) > 3000:
                if current_page:
                    pages.append("\n\n".join(current_page))
                current_page = [item]
                current_length = len(item)
            else:
                current_page.append(item)
                current_length += len(item)
        
        # Add the last page if it has content
        if current_page:
            pages.append("\n\n".join(current_page))
        
        # Convert each page of text to an image
        images = []
        for page_text in pages:
            # Create a new image with white background and the text content
            img = create_text_image(page_text, max_width, font_size, margin)
            
            # Convert the image to a BytesIO object for Discord upload
            img_byte_arr = BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)  # Reset file pointer to beginning
            
            images.append(img_byte_arr)
        
        return images
    except Exception as e:
        # Log any errors that occur during conversion
        print(f"Error converting DOCX: {e}")
        return None

def create_text_image(text, max_width=1024, font_size=16, margin=20, watermark_text="GridZer0 Bot", watermark_opacity=0.3):
    """Create an image from text with watermark.
    
    This function renders text onto an image with a white background and adds
    a watermark. It handles text wrapping and paragraph spacing.
    
    Args:
        text: The text to render on the image
        max_width: Maximum width of the image in pixels (default: 1024)
        font_size: Font size to use for text (default: 16)
        margin: Margin around text in pixels (default: 20)
        watermark_text: Text to use as watermark (default: "GridZer0 Bot")
        watermark_opacity: Opacity level for the watermark (default: 0.3)
        
    Returns:
        A PIL Image object containing the rendered text
    """
    try:
        # Try to load a nice font, fall back to default if not available
        try:
            # Try to use Arial font first (common on Windows)
            font = ImageFont.truetype("arial.ttf", font_size)
            watermark_font = ImageFont.truetype("arial.ttf", int(font_size * 2))  # Larger font for watermark
        except:
            try:
                # Try DejaVu Sans font (common on Linux)
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size)
                watermark_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", int(font_size * 2))
            except:
                # If all else fails, use the default font
                font = ImageFont.load_default()
                watermark_font = ImageFont.load_default()
        
        # Calculate line wrapping
        lines = []
        line_height = font_size + 4  # Add a little extra space between lines
        
        # Wrap text to fit within max_width
        # Process each paragraph separately to maintain paragraph breaks
        for paragraph in text.split('\n'):
            if paragraph.strip():  # Skip empty paragraphs
                # Calculate how many characters can fit on a line based on font size and width
                wrapped_lines = textwrap.wrap(paragraph, width=max_width//font_size*2)
                lines.extend(wrapped_lines)
                lines.append('')  # Add empty line after each paragraph
        
        # If the last line is empty, remove it
        if lines and not lines[-1]:
            lines = lines[:-1]
        
        # Calculate image height based on number of lines
        height = (len(lines) * line_height) + (margin * 2)
        
        # Create a new image with white background
        image = Image.new("RGB", (max_width, height), color=(255, 255, 255))
        draw = ImageDraw.Draw(image)
        
        # Draw the text on the image
        y_position = margin
        for line in lines:
            draw.text((margin, y_position), line, font=font, fill=(0, 0, 0))
            y_position += line_height
        
        # Add watermark - very simple approach
        # Create a semi-transparent text directly on the image
        draw = ImageDraw.Draw(image)
        
        # Use a light gray color with the specified opacity
        gray_level = int(200 * (1 - watermark_opacity) + 55)  # Ranges from 55-255 based on opacity
        
        # Draw the watermark text in the center
        # Calculate the center position
        center_x = max_width // 2
        center_y = height // 2
        
        # Draw the watermark text
        draw.text(
            (center_x, center_y),
            watermark_text,
            font=watermark_font,
            fill=(gray_level, gray_level, gray_level),
            anchor="mm"  # Center alignment (middle-middle)
        )
        
        return image
    except Exception as e:
        # Log any errors and return a simple error image
        print(f"Error creating text image: {e}")
        # Return a small image with error message
        image = Image.new("RGB", (400, 100), color=(255, 255, 255))
        draw = ImageDraw.Draw(image)
        draw.text((10, 10), f"Error rendering text: {str(e)}", fill=(255, 0, 0))
        return image

class DOCXOptionsView(View):
    """A view class to present DOCX handling options as buttons.
    
    This class creates an interactive UI with buttons that allow users to choose
    how they want to view the DOCX file - either in a new thread or in the current channel.
    """
    def __init__(self, message, attachment, processing_msg):
        # No timeout to keep buttons active indefinitely
        super().__init__(timeout=None)
        self.message = message  # The original message containing the DOCX
        self.attachment = attachment  # The DOCX file attachment
        self.processing_msg = processing_msg  # Message showing processing status
        self.button_clicked = False  # Flag to track if a button has been clicked

    @discord.ui.button(label="Create Thread", style=discord.ButtonStyle.primary)
    async def create_thread(self, interaction: discord.Interaction, button: Button):
        """Handle the 'Create Thread' button press.
        
        Creates a new thread for the DOCX and posts each page as an image.
        
        Args:
            interaction: The Discord interaction object
            button: The button that was pressed
        """
        # Prevent multiple clicks from processing the same DOCX multiple times
        if self.button_clicked:
            await interaction.response.defer()
            return
        self.button_clicked = True
        
        # Disable all buttons to prevent further interaction
        for item in self.children:
            item.disabled = True
        
        # Update the message with disabled buttons
        await interaction.response.edit_message(view=self)
        # Process the DOCX in a thread
        await self.process_docx(use_thread=True)

    @discord.ui.button(label="Post Here", style=discord.ButtonStyle.secondary)
    async def post_here(self, interaction: discord.Interaction, button: Button):
        """Handle the 'Post Here' button press.
        
        Posts all DOCX pages as images in the current channel.
        
        Args:
            interaction: The Discord interaction object
            button: The button that was pressed
        """
        # Prevent multiple clicks from processing the same DOCX multiple times
        if self.button_clicked:
            await interaction.response.defer()
            return
        self.button_clicked = True
        
        # Disable all buttons to prevent further interaction
        for item in self.children:
            item.disabled = True
        
        # Update the message with disabled buttons
        await interaction.response.edit_message(view=self)
        # Process the DOCX in the current channel
        await self.process_docx(use_thread=False)

    async def process_docx(self, use_thread):
        """Process the DOCX, either in a thread or the current channel.
        
        This is the main processing function that handles the DOCX conversion
        and posting of images to the appropriate location.
        
        Args:
            use_thread: Boolean indicating whether to create a new thread
        """
        try:
            # Update processing message and remove buttons
            await self.processing_msg.edit(content="Converting DOCX to images...", view=None)
            
            # Convert DOCX to images
            docx_bytes = await self.attachment.read()  # Read the DOCX file data
            images = await convert_docx_to_images(docx_bytes)  # Convert to images
            if not images:
                raise ValueError("Failed to convert DOCX to images or document is empty")

            # Determine target channel (thread or current channel)
            if use_thread:
                # Create thread with simplified approach
                thread_name = self.attachment.filename  # Use DOCX filename as thread name
                # Truncate thread name if too long (Discord limit)
                if len(thread_name) > 100:
                    thread_name = thread_name[:97] + "..."
                
                try:
                    # Create thread with retry mechanism for rate limits
                    target = await self.create_thread_with_retry(thread_name)
                    if not target:
                        # If thread creation failed, fall back to current channel
                        target = self.message.channel
                        final_message = "⚠️ Could not create thread. Posted DOCX in current channel instead."
                    else:
                        final_message = f"{target.mention}"  # Mention the thread
                except Exception as e:
                    print(f"Thread creation error: {e}")
                    target = self.message.channel
                    final_message = "⚠️ Could not create thread. Posted DOCX in current channel instead."
            else:
                # Use current channel
                target = self.message.channel
                final_message = None

            # Post images to the target (thread or channel)
            await self.post_images(target, images, use_thread)
            
            # Update or delete processing message
            if final_message:
                await self.processing_msg.edit(content=final_message)
            else:
                await self.processing_msg.delete()
            
            # Delete original message with DOCX attachment to keep channel clean
            try:
                await self.message.delete()
            except Forbidden:
                print("Bot lacks permission to delete messages.")
            except Exception as e:
                print(f"Error deleting original message: {e}")
                
        except Forbidden as e:
            # Handle permission errors
            error_msg = f"⚠️ Permission error: {str(e)}"
            await self.processing_msg.edit(content=error_msg)
            print(f"Discord permission error: {e}")
        except Exception as e:
            # Handle general errors
            error_msg = f"❌ Error processing DOCX: {str(e)}"
            await self.processing_msg.edit(content=error_msg)
            print(f"Error processing DOCX: {e}")
            # Clean up thread if it was created but processing failed
            if use_thread and 'target' in locals() and isinstance(target, discord.Thread):
                try:
                    await target.delete()
                except Exception:
                    pass
    
    async def create_thread_with_retry(self, thread_name):
        """Create a thread with retry logic for rate limits.
        
        Discord has rate limits that can cause thread creation to fail.
        This function implements retry logic with exponential backoff.
        
        Args:
            thread_name: Name for the new thread
            
        Returns:
            The created thread object or None if creation failed
        """
        max_retries = 3  # Maximum number of attempts
        retry_delay = 2  # Start with 2 seconds delay
        
        for attempt in range(max_retries):
            try:
                # Create thread directly using channel's create_thread method
                thread = await self.message.channel.create_thread(
                    name=thread_name,
                    type=discord.ChannelType.public_thread,
                    auto_archive_duration=1440  # 24 hours
                )
                
                # Wait a bit longer for the notification message to appear
                await asyncio.sleep(2)
                
                # Delete the thread creation notification to keep channel clean
                try:
                    # First try to find system messages about thread creation
                    found = False
                    async for msg in self.message.channel.history(limit=10):
                        # Check for system message type
                        if msg.type == discord.MessageType.thread_created:
                            await msg.delete()
                            found = True
                            break
                        
                        # Fallback: Check for content that looks like a thread notification
                        # This handles cases where the message type check doesn't work
                        if msg.author.bot and "started a thread" in msg.content and thread_name in msg.content:
                            await msg.delete()
                            found = True
                            break
                    
                    if not found:
                        print(f"Could not find thread notification message to delete")
                except Exception as e:
                    print(f"Error deleting thread notification: {e}")
                
                return thread
                
            except discord.Forbidden as e:
                # Handle permission errors during thread creation
                print(f"Thread creation attempt {attempt+1} failed due to permissions: {e}")
                if attempt == max_retries - 1:
                    return None  # Failed after all retries
                await asyncio.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
            except Exception as e:
                # Handle other errors during thread creation
                print(f"Unexpected error in thread creation: {e}")
                return None
        
        return None  # Failed to create thread after all retries
    
    async def post_images(self, target, images, use_thread):
        """Post images to the target channel or thread.
        
        Args:
            target: The Discord channel or thread to post images to
            images: List of BytesIO objects containing the images
            use_thread: Boolean indicating whether we're using a thread
        """
        try:
            # Determine starting index (skip first page if already posted in thread)
            # This is because the first page is often shown in the thread creation notification
            start_idx = 1 if use_thread and isinstance(target, discord.Thread) else 0
            
            # If using thread and we have the first image, post it first
            if use_thread and isinstance(target, discord.Thread) and images:
                images[0].seek(0)  # Reset file pointer to beginning
                await target.send(file=discord.File(fp=images[0], filename="page_1.png"))
            
            # Post remaining images
            for i, img in enumerate(images[start_idx:], start_idx + 1):
                img.seek(0)  # Reset file pointer to beginning
                await target.send(file=discord.File(fp=img, filename=f"page_{i}.png"))
                await asyncio.sleep(0.5)  # Small delay to avoid rate limits but still be safe
                
        except Exception as e:
            # Propagate any errors during posting
            raise Exception(f"Error posting images: {e}")

async def handle_docx(message):
    """Main handler function for DOCX attachments, offering processing options.
    
    This is the entry point function called by the bot when a DOCX is detected.
    
    Args:
        message: The Discord message containing the DOCX attachment
        
    Returns:
        Boolean indicating whether the message was handled as a DOCX
    """
    # Check if the message has any attachments
    if not message.attachments:
        return False
    
    # Find DOCX attachments by checking file extensions
    docx_attachments = [att for att in message.attachments if att.filename.lower().endswith('.docx')]
    if not docx_attachments:
        return False  # No DOCX attachments found
    
    # Process only the first DOCX attachment if multiple are present
    attachment = docx_attachments[0]
    
    # Check file size to prevent processing very large files
    if attachment.size > 8 * 1024 * 1024:  # 8MB limit
        await message.channel.send(f"⚠️ DOCX too large: {attachment.filename} (max 8MB)")
        return True
    
    # Check basic permissions to ensure the bot can function
    channel = message.channel
    bot_member = channel.guild.me if hasattr(channel, 'guild') else None
    
    if bot_member:
        permissions = channel.permissions_for(bot_member)
        missing_permissions = []
        
        # Check for required permissions
        if not permissions.send_messages:
            missing_permissions.append("Send Messages")
        if not permissions.attach_files:
            missing_permissions.append("Attach Files")
        
        # Notify if permissions are missing
        if missing_permissions:
            await message.channel.send(
                f"⚠️ Bot lacks required permissions: {', '.join(missing_permissions)}. "
                f"Please ensure the bot has these permissions in this channel."
            )
            return True
    
    try:
        # Present options to the user via buttons
        processing_msg = await message.channel.send("Choose an option for this DOCX:")
        view = DOCXOptionsView(message, attachment, processing_msg)
        await processing_msg.edit(view=view)
        return True  # Successfully handled the DOCX
    except Exception as e:
        # Handle any errors during setup
        print(f"Error presenting DOCX options: {e}")
        return False  # Failed to handle the DOCX